#!/usr/bin/env python3
"""Convert custom XML (indexing format) to DOCX."""

import xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import sys
import os

INPUT = os.path.join(os.path.dirname(__file__), "Нежкин_ДС_ВКР_промежуточный_с_пометками.xml")
OUTPUT = os.path.join(os.path.dirname(__file__), "Нежкин_ДС_ВКР_промежуточный_с_пометками.docx")


def is_heading(text):
    """Detect chapter/section headings."""
    text_stripped = text.strip()
    # Main structural headings (all caps)
    upper_headings = [
        "ВВЕДЕНИЕ", "ЗАКЛЮЧЕНИЕ", "СОДЕРЖАНИЕ",
        "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ",
        "ПРИЛОЖЕНИЕ", "РЕФЕРАТ",
    ]
    for h in upper_headings:
        if text_stripped.upper().startswith(h):
            return 1  # Heading level 1

    # Chapter headings: "1 Обзор предметной области"
    if re.match(r'^\d+\s+[А-ЯA-Z]', text_stripped):
        return 1
    # Section headings: "1.1 Анализ..."
    if re.match(r'^\d+\.\d+\s+', text_stripped):
        return 2
    return 0


def is_title_page(index):
    """Title page elements (before СОДЕРЖАНИЕ)."""
    return index < 64


def main():
    tree = ET.parse(INPUT)
    root = tree.getroot()

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    pf = style.paragraph_format
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.5

    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)

    # Collect table parent indices to identify table cells
    table_indices = set()
    for obj in root.findall('.//object[@object_type="table"]'):
        table_indices.add(obj.get('index'))

    # Track section (TOC) indices
    section_indices = set()
    for obj in root.findall('.//object[@object_type="section"]'):
        section_indices.add(obj.get('index'))

    for elem in root:
        if elem.tag == 'object':
            continue

        if elem.tag == 'paragraph':
            text = (elem.text or '').strip()
            if not text:
                continue

            index_str = elem.get('index', '0')
            try:
                index = int(index_str)
            except ValueError:
                index = 9999  # non-numeric index, treat as body content
            parent = elem.get('parent_index')

            # Skip table cells - they lose structure in plain paragraphs
            # But include them as indented text
            heading_level = is_heading(text)

            if heading_level and not parent:
                p = doc.add_heading(text, level=heading_level)
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)
                    run.font.color.rgb = None  # Remove blue color from headings
            else:
                p = doc.add_paragraph(text)
                if is_title_page(index):
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if parent and parent in table_indices:
                    # Table cell content - indent slightly
                    p.paragraph_format.left_indent = Cm(1)
                    p.runs[0].font.size = Pt(12) if p.runs else None

    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")
    print(f"Paragraphs in document: {len(doc.paragraphs)}")


if __name__ == '__main__':
    main()
